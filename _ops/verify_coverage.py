"""
Verify coverage: extract list of CONCEPTS từ 5 file, check xem mỗi concept đã
được mention trong file gốc (đã annotate) chưa.
"""
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = r'C:\xampp\htdocs\Web-chong-hang-gia-main'
ORIGINAL = os.path.join(ROOT, 'TAI_LIEU_NGHIEP_VU_VNTRUST_FULL.docx')
FILES_5 = [
    '1. hệ thống dữ liệu sản phẩm doanh nghiệp.docx',
    '2.hệ thống phản ảnh sản phẩm hàng hóa từ khách hàng là người dùng.docx',
    '3. hệ thống chứa các thuật toán để kiểm tra sản phẩm hàng hóa.docx',
    '4. hệ thống truy vấn dữ liệu và kết xuất báo cáo.docx',
    '5. Giao diện chung của toàn bộ nền tảng.docx',
]

# Read all annotated doc text
orig = Document(ORIGINAL)
orig_text = '\n'.join(p.text for p in orig.paragraphs)
for t in orig.tables:
    for r in t.rows:
        for c in r.cells:
            orig_text += '\n' + c.text

print(f'Original (annotated) doc: {len(orig_text)} chars\n')

# For each of 5 files, extract key concepts (unique phrases) and check coverage
for f in FILES_5:
    path = os.path.join(ROOT, f)
    d = Document(path)
    concepts = set()

    # From paragraphs — extract bold/heading-like text or distinctive phrases
    for p in d.paragraphs:
        txt = p.text.strip()
        if not txt: continue
        # Skip very short or generic
        if len(txt) < 8: continue
        # Headings/sections - high signal
        is_heading = (
            p.style.name.startswith('Heading') if p.style else False
        ) or re.match(r'^\d+\.\s|\b(MODULE|PHÂN HỆ|UC\d|BƯỚC|STEP|API)\b', txt, re.I)
        if is_heading:
            concepts.add(txt[:80])

    # From tables — each row often has concept
    for t in d.tables:
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells if c.text.strip()]
            if cells and len(cells[0]) > 3 and len(cells[0]) < 60:
                concepts.add(cells[0])

    print(f'═══ {f} ═══')
    print(f'  Total concepts extracted: {len(concepts)}')

    # Check coverage
    covered = 0
    not_covered = []
    for c in concepts:
        # Normalize whitespace
        c_norm = re.sub(r'\s+', ' ', c).strip()
        if len(c_norm) < 5: continue
        # Check if first 30 chars appear in original
        snippet = c_norm[:30].lower()
        if snippet in orig_text.lower():
            covered += 1
        else:
            not_covered.append(c_norm[:70])

    pct = round(covered / len(concepts) * 100) if concepts else 0
    print(f'  ✅ Covered in original: {covered}/{len(concepts)} ({pct}%)')

    if not_covered:
        print(f'  ❌ NOT FOUND ({len(not_covered)}):')
        for nc in not_covered[:30]:
            print(f'      • {nc}')
        if len(not_covered) > 30:
            print(f'      ... +{len(not_covered)-30} khác')
    print()
