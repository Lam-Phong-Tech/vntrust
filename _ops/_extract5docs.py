import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = r'C:\xampp\htdocs\Web-chong-hang-gia-main'
FILES = [
    '1. hệ thống dữ liệu sản phẩm doanh nghiệp.docx',
    '2.hệ thống phản ảnh sản phẩm hàng hóa từ khách hàng là người dùng.docx',
    '3. hệ thống chứa các thuật toán để kiểm tra sản phẩm hàng hóa.docx',
    '4. hệ thống truy vấn dữ liệu và kết xuất báo cáo.docx',
    '5. Giao diện chung của toàn bộ nền tảng.docx',
]

OUT = os.path.join(ROOT, '_ops', '_5docs_combined.txt')
with open(OUT, 'w', encoding='utf-8') as out:
    for f in FILES:
        path = os.path.join(ROOT, f)
        if not os.path.exists(path):
            out.write(f'\n\n===== MISSING: {f} =====\n')
            continue
        d = Document(path)
        out.write(f'\n\n========================================\n')
        out.write(f'FILE: {f}\n')
        out.write(f'========================================\n\n')
        for p in d.paragraphs:
            if p.text.strip():
                out.write(p.text + '\n')
        # Also tables
        for i, t in enumerate(d.tables):
            out.write(f'\n--- Table {i} ({len(t.rows)} rows × {len(t.columns)} cols) ---\n')
            for r in t.rows:
                cells = [c.text.replace(chr(10), ' / ').strip()[:200] for c in r.cells]
                out.write(' | '.join(cells) + '\n')

print(f'Wrote {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
