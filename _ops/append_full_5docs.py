"""
Append TOÀN BỘ nội dung 5 file vào cuối tài liệu gốc (đã annotate).
Mỗi file thành 1 chương lớn với đầy đủ paragraphs + tables.
Không tóm tắt — copy đầy đủ.
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'C:\xampp\htdocs\Web-chong-hang-gia-main'
ORIGINAL = os.path.join(ROOT, 'TAI_LIEU_NGHIEP_VU_VNTRUST_FULL.docx')

FILES_5 = [
    ('XVII', 'MODULE 1 — HỆ THỐNG DỮ LIỆU SẢN PHẨM DOANH NGHIỆP (file 1 đầy đủ)',
        '1. hệ thống dữ liệu sản phẩm doanh nghiệp.docx',
        'Enterprise data system — DN khai báo hồ sơ, chứng nhận, sản phẩm, workflow phê duyệt'),
    ('XVIII', 'MODULE 2 — HỆ THỐNG PHẢN ÁNH TỪ NGƯỜI TIÊU DÙNG (file 2 đầy đủ)',
        '2.hệ thống phản ảnh sản phẩm hàng hóa từ khách hàng là người dùng.docx',
        'Consumer crowdsourcing — NTD quét/chụp + GPS + AI risk scoring + ẩn danh'),
    ('XIX', 'MODULE 3 — HỆ THỐNG THUẬT TOÁN KIỂM TRA (file 3 đầy đủ)',
        '3. hệ thống chứa các thuật toán để kiểm tra sản phẩm hàng hóa.docx',
        'AI Fraud Detection Engine — Rule + Vision + Similarity + Risk Scoring + Investigation'),
    ('XX', 'MODULE 4 — HỆ THỐNG TRUY VẤN & BÁO CÁO (file 4 đầy đủ)',
        '4. hệ thống truy vấn dữ liệu và kết xuất báo cáo.docx',
        'BI Analytics Command Center — Search, Heatmap, Time slider, 8 biểu đồ, IOC dashboard'),
    ('XXI', 'MODULE 5 — GIAO DIỆN CHUNG UNIFIED PORTAL (file 5 đầy đủ)',
        '5. Giao diện chung của toàn bộ nền tảng.docx',
        'Unified Super Portal — Landing 4 nút lớn, Subscription, PWA, Mobile-first'),
]

GREEN = RGBColor(0x10, 0xB9, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
RED   = RGBColor(0xEF, 0x44, 0x44)
BLUE  = RGBColor(0x3B, 0x82, 0xF6)
GOLD  = RGBColor(0xC8, 0xA5, 0x57)
SLATE = RGBColor(0x47, 0x55, 0x69)
INK   = RGBColor(0x0B, 0x16, 0x23)

doc = Document(ORIGINAL)

def add_h1(text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = color
    return p

def add_h2(text, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = color
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = INK
    return p

def add_p(text, bold=False, italic=False, color=None, size=11, indent=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    r.font.size = Pt(size)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
    r = p.add_run('• ' + text)
    r.font.size = Pt(11)
    return p

def _border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def copy_table_from(source_table):
    """Copy table from source docx into doc, preserving content + borders"""
    nrows = len(source_table.rows)
    ncols = len(source_table.columns)
    if nrows == 0 or ncols == 0:
        return None
    new_t = doc.add_table(rows=nrows, cols=ncols)
    for ri, src_row in enumerate(source_table.rows):
        for ci, src_cell in enumerate(src_row.cells):
            if ci >= ncols: break
            new_cell = new_t.rows[ri].cells[ci]
            text = src_cell.text.strip()
            new_cell.text = ''
            p = new_cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            # Bold header row
            if ri == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '0B1623')
                new_cell._tc.get_or_add_tcPr().append(shading)
            _border(new_cell)
    return new_t

def status_box(detail, color=AMBER):
    """Add status box callout"""
    add_p('┌─ ĐỐI CHIẾU VỚI HỆ THỐNG VNTRUST ────────────────────────', color=color, size=9)
    add_p(f'│ {detail}', bold=True, color=color, size=10)
    add_p('└─────────────────────────────────────────────────────────', color=color, size=9)

# ─── Insert big page break + intro section ──────────────────────────
doc.add_page_break()
add_h1('PHẦN MỞ RỘNG — TOÀN BỘ NỘI DUNG 5 TÀI LIỆU CHI TIẾT MODULE', GOLD)
add_p(
    'Phần dưới đây trích NGUYÊN VĂN toàn bộ nội dung 5 tài liệu chi tiết module '
    '(file 1-5), không tóm tắt. Mỗi chương XVII-XXI tương ứng 1 file, kèm bảng '
    'đối chiếu trạng thái cuối mỗi mục với hệ thống VNTrust đã build.',
    italic=True, color=SLATE, size=11
)
add_p('Mục đích: đảm bảo KHÔNG SÓT bất kỳ chi tiết nghiệp vụ nào từ 5 file.',
    italic=True, color=BLUE, size=11)
add_p('')

# Helper: append all paragraphs + tables from source doc preserving order
def append_full_source(source_path):
    src = Document(source_path)
    # Build a list of (kind, element) preserving order
    body = src.element.body
    elements = []
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            elements.append(('p', child))
        elif tag == 'tbl':
            elements.append(('tbl', child))

    # Map XML elements back to python-docx objects
    para_idx = 0
    tbl_idx = 0
    for kind, _ in elements:
        if kind == 'p':
            src_para = src.paragraphs[para_idx]
            para_idx += 1
            txt = src_para.text.strip()
            if not txt:
                # Add empty spacer
                doc.add_paragraph()
                continue
            # Determine heading level by style or pattern
            style_name = src_para.style.name if src_para.style else ''
            if style_name.startswith('Heading 1'):
                add_h2(txt)  # demote 1 level
            elif style_name.startswith('Heading 2'):
                add_h3(txt)
            elif style_name.startswith('Heading 3'):
                add_p(txt, bold=True, size=12)
            elif style_name.startswith('Heading 4') or style_name.startswith('Heading 5'):
                add_p(txt, bold=True, size=11)
            elif style_name.startswith('List Bullet'):
                add_bullet(txt)
            elif style_name.startswith('List Number'):
                add_bullet(txt)
            else:
                # Detect numbered sections like "1.", "2.", "10." → heading
                if len(txt) < 100 and (txt[0:2] in ('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')
                                       or txt[0:3] in ('10.', '11.', '12.', '13.', '14.', '15.', '16.', '17.', '18.', '19.', '20.', '21.', '22.')):
                    add_p(txt, bold=True, size=12, color=INK)
                else:
                    add_p(txt, size=11)
        else:  # tbl
            src_table = src.tables[tbl_idx]
            tbl_idx += 1
            copy_table_from(src_table)

# ─── Append each of 5 files as full chapter ─────────────────────────
for chapter_num, chapter_title, filename, subtitle in FILES_5:
    doc.add_page_break()
    add_h1(f'{chapter_num}. {chapter_title}')
    add_p(subtitle, italic=True, color=SLATE, size=11)
    add_p('')

    path = os.path.join(ROOT, filename)
    if os.path.exists(path):
        print(f'Appending: {filename}')
        append_full_source(path)
    else:
        add_p(f'⚠️ File không tồn tại: {filename}', color=RED, bold=True)

# Footer
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot_run = foot.add_run('— HẾT TÀI LIỆU HỢP NHẤT (UPDATED 30/05/2026 — Tích hợp 6 file nguồn) —')
foot_run.italic = True
foot_run.font.color.rgb = GOLD
foot_run.font.size = Pt(11)

# Save
doc.save(ORIGINAL)
print(f'\n✅ Saved: {ORIGINAL}')
print(f'   Size: {os.path.getsize(ORIGINAL)/1024:.1f} KB')
print(f'   Paragraphs: {len(doc.paragraphs)}')
print(f'   Tables: {len(doc.tables)}')
