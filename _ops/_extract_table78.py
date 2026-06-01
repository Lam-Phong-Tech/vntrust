import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
d = Document(r'C:\xampp\htdocs\Web-chong-hang-gia-main\TAI_LIEU_NGHIEP_VU_VNTRUST_FULL.docx')
# Show tables around permission matrix
for i in [77, 78, 79, 110, 111, 112, 113]:
    if i >= len(d.tables): continue
    t = d.tables[i]
    print(f'=== Table {i} ({len(t.rows)} rows x {len(t.columns)} cols) ===')
    for r in t.rows:
        # Show raw repr to see hidden chars
        cells = [repr(c.text.strip()) for c in r.cells]
        print(' | '.join(cells))
    print()
