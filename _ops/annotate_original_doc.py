"""
Chỉnh sửa file gốc TAI_LIEU_NGHIEP_VU_VNTRUST_FULL.docx:
- Giữ nguyên 100% nội dung gốc
- Thêm status badge sau mỗi section quan trọng
- Bổ sung sections mới từ 5 file chi tiết (Unified Portal, Investigation Console, ...)
- Thêm Appendix chi tiết gap analysis
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = r'C:\xampp\htdocs\Web-chong-hang-gia-main'
FILE = os.path.join(ROOT, 'TAI_LIEU_NGHIEP_VU_VNTRUST_FULL.docx')

# Colors
GREEN = RGBColor(0x10, 0xB9, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
RED   = RGBColor(0xEF, 0x44, 0x44)
BLUE  = RGBColor(0x3B, 0x82, 0xF6)
GOLD  = RGBColor(0xC8, 0xA5, 0x57)
SLATE = RGBColor(0x47, 0x55, 0x69)

doc = Document(FILE)

# ─── Helpers — insert paragraph AFTER a reference paragraph ──────────
def insert_para_after(ref_para, text, bold=False, italic=False, color=None, size=None, style=None):
    """Insert new paragraph immediately after ref_para. Returns new paragraph for chaining."""
    new_p = OxmlElement('w:p')
    ref_para._p.addnext(new_p)
    new_para = Paragraph(new_p, ref_para._parent)
    if style:
        try:
            new_para.style = doc.styles[style]
        except KeyError:
            pass
    run = new_para.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    if size: run.font.size = Pt(size)
    return new_para

def insert_status_box(ref_para, status, detail, missing=None, bonus=None):
    """Insert a status callout box after ref_para. Returns last inserted para."""
    icon = {'done': '✅', 'partial': '🟡', 'missing': '❌', 'bonus': '✨'}.get(status, '•')
    color = {'done': GREEN, 'partial': AMBER, 'missing': RED, 'bonus': BLUE}.get(status, SLATE)

    last = ref_para
    # Status line
    last = insert_para_after(last, f'  ┌─ TRẠNG THÁI HỆ THỐNG ─────────────────────────────────', color=color, size=9)
    last = insert_para_after(last, f'  │ {icon} {detail}', bold=True, color=color, size=10)
    if missing:
        last = insert_para_after(last, f'  │ 🔴 Thiếu: {missing}', color=RED, size=9, italic=True)
    if bonus:
        last = insert_para_after(last, f'  │ ✨ Bonus đã build: {bonus}', color=BLUE, size=9, italic=True)
    last = insert_para_after(last, f'  └────────────────────────────────────────────────────────', color=color, size=9)
    last = insert_para_after(last, '', size=6)  # spacer
    return last

# ─── Find paragraphs by text contains ────────────────────────────────
def find_para(text_contains, after_idx=0):
    """Find first paragraph whose text contains substring, starting from index."""
    for i, p in enumerate(doc.paragraphs):
        if i < after_idx: continue
        if text_contains in p.text:
            return i, p
    return -1, None

print('=== STEP 1: Add status legend at top ===')
# Find paragraph "MỤC LỤC" (the second one — section divider)
for i, p in enumerate(doc.paragraphs):
    if 'MỤC LỤC' in p.text and i > 5 and i < 25:
        # Insert legend right after MỤC LỤC heading
        ref = p
        ref = insert_para_after(ref, '', size=6)
        ref = insert_para_after(ref, '═══ BẢNG MÀU TRẠNG THÁI (cập nhật 30/05/2026) ═══', bold=True, color=GOLD, size=11)
        ref = insert_para_after(ref, '  ✅ ĐÃ CÓ — Tính năng đã build + deploy production', color=GREEN, size=10)
        ref = insert_para_after(ref, '  🟡 MỘT PHẦN — Có backend nhưng UI sơ sài, hoặc thiếu sub-features', color=AMBER, size=10)
        ref = insert_para_after(ref, '  ❌ CHƯA CÓ — Cần build từ đầu', color=RED, size=10)
        ref = insert_para_after(ref, '  ✨ BONUS — Đã build vượt yêu cầu tài liệu', color=BLUE, size=10)
        ref = insert_para_after(ref, '', size=6)
        ref = insert_para_after(ref,
            'Tài liệu này được CẬP NHẬT từ file gốc (V2.0 14/05/2026) — bổ sung trạng thái '
            'sau khi đối chiếu hệ thống VNTrust đã build (44 task hoàn thành) + tích hợp '
            '5 tài liệu chi tiết module (1-5).',
            italic=True, color=SLATE, size=10)
        ref = insert_para_after(ref, '', size=6)
        break

print('=== STEP 2: Annotate major sections ===')

# Helper: annotate by finding heading text
def annotate(heading_text, after_idx, status, detail, missing=None, bonus=None):
    idx, p = find_para(heading_text, after_idx)
    if p:
        insert_status_box(p, status, detail, missing, bonus)
        print(f'  ✓ Annotated: "{heading_text[:60]}"')
        return idx
    print(f'  ✗ NOT FOUND: "{heading_text}"')
    return after_idx

last_idx = 0
# I. HỆ THỐNG VNTRUST → 6 phân hệ
last_idx = annotate('PHÂN HỆ 1 — Quản lý hồ sơ doanh nghiệp', last_idx, 'done',
    'ĐÃ CÓ — KYC flow + Quản lý SP/Lô/Chứng nhận + Sub-role UC03',
    missing='Model Nhà máy + Dây chuyền + Logo/Thương hiệu riêng + ESG/CSR fields',
    bonus='Profile expansion: avatar + địa chỉ + ngày sinh + CCCD + giới tính')

last_idx = annotate('PHÂN HỆ 2 — Tạo và quản lý mã định danh', last_idx, 'done',
    'ĐÃ CÓ ĐẦY ĐỦ — UID encrypted + Bulk import + Print export',
    bonus='QR zoom modal mobile 320px + level H + scanner fps 15')

last_idx = annotate('PHÂN HỆ 3 — Xác thực sản phẩm', last_idx, 'done',
    'ĐÃ CÓ ĐẦY ĐỦ — Quét QR + Verify + Báo cáo nghi vấn + AI Photo Vision',
    bonus='AI Photo Verify với MobileNet client-side + cosine similarity 2 vector')

last_idx = annotate('PHÂN HỆ 4 — Giám sát', last_idx, 'done',
    'ĐÃ CÓ — Dashboard live stats + Anomaly cron + 4 heatmap layers VietMap',
    bonus='Time slider animation + 4 layer toggle + GPS marker compass cone')

last_idx = annotate('PHÂN HỆ 5 — Quản lý cảnh báo', last_idx, 'done',
    'ĐÃ CÓ — Cron escalate/digest + Email SMTP + 4 báo cáo định kỳ',
    bonus='Lifecycle config UI + tùy chỉnh ngành (Sữa 60d, Mỹ phẩm 90d)')

last_idx = annotate('PHÂN HỆ 6 — Tích hợp', last_idx, 'partial',
    'MỘT PHẦN — Có Webhook ERP + API public + Shopee stub',
    missing='TikTok Shop integration + Apache Superset + Kafka stream + OpenAPI docs')

# II. TRUSTCHECK
last_idx = annotate('II.  NỀN TẢNG GIÁM SÁT VÀ PHẢN ÁNH', last_idx, 'done',
    'ĐÃ CÓ ĐẦY ĐỦ 5 phân hệ — Tra cứu + Phản ánh ẩn danh + Admin + Report + API',
    missing='Wizard step-by-step 4 bước upload + Reward gamification + Livestream Detection')

# III. BẢO MẬT
last_idx = annotate('III.  THIẾT KẾ HỆ THỐNG TIẾP NHẬN', last_idx, 'partial',
    'ĐÃ CÓ 85% — 3 vault (Identity/Report/AnonSession) + AES-256 + OTP + Audit log',
    missing='Unmask PII flow UI (yêu cầu giải mã từ cơ quan chức năng) + microservices thật')

# IV. ĐỐI SÁT
last_idx = annotate('IV.  HỆ THỐNG ĐỐI SÁT', last_idx, 'partial',
    'MỘT PHẦN 60% — Trust Score 0-100 + Anomaly cron + Client MobileNet cosine sim',
    missing='Server-side YOLOv8 + Siamese CNN + GAN Detection + PaddleOCR + Investigation Console + Side-by-side image diff + 6-color checklist UI')

# V. CẢNH BÁO VÒNG ĐỜI
last_idx = annotate('V.  CẢNH BÁO VÒNG ĐỜI', last_idx, 'done',
    'ĐÃ CÓ ĐẦY ĐỦ — 3 lớp cảnh báo + Multi-channel email + Cron lifecycle-check',
    bonus='UI lifecycle-config cho DN tự chỉnh threshold per ngành hàng')

# VI. BẢNG ĐIỀU KHIỂN TUÂN THỦ
last_idx = annotate('VI.  BẢNG ĐIỀU KHIỂN TUÂN THỦ', last_idx, 'done',
    'ĐÃ CÓ ĐẦY ĐỦ — Checklist động + Priority + Heatmap khu vực (VietMap)',
    bonus='Time slider animation cho heatmap evolve theo thời gian')

# VII. HẬU KIỂM
last_idx = annotate('VII. QUẢN LÝ DỮ LIỆU HẬU KIỂM', last_idx, 'done',
    'ĐÃ CÓ ĐẦY ĐỦ — Rule engine 25+ tiêu chuẩn + Auto compare + 3 đối tượng upload')

# X. KYC + Phân quyền
last_idx = annotate('PHÂN QUYỀN & XÁC THỰC NGƯỜI DÙNG', last_idx, 'done',
    'ĐÃ CÓ ĐẦY ĐỦ — 4 platform role + UC03 sub-role (company_admin/staff_input/warehouse/viewer)',
    missing='Cán bộ tư vấn + Cơ quan chức năng roles + Subscription tiers (Basic/Pro/Enterprise/Gov)')

print('=== STEP 3: Annotate UC01-UC20 inline ===')
UC_STATUS = {
    'UC01': ('done', 'KYC registration flow'),
    'UC02': ('done', 'KYC document verify'),
    'UC03': ('done', 'Sub-role nội bộ DN (company_admin/staff_input/warehouse/viewer) — vừa build Sprint UC03'),
    'UC04': ('done', 'Inventory page'),
    'UC05': ('done', 'Inventory lô hàng'),
    'UC06': ('done', 'Certificates page + upload'),
    'UC07': ('done', 'Bulk UID generation'),
    'UC08': ('done', 'QR export page'),
    'UC09': ('done', 'Verify page + scan'),
    'UC10': ('done', 'Result page + 3D image'),
    'UC11': ('done', 'Report submission flow'),
    'UC12': ('done', 'Dashboard + 4 heatmap layers'),
    'UC13': ('done', 'VietMap với GPS + heatmap + time slider'),
    'UC14': ('done', 'System-config UI (Sprint Bundle 1)'),
    'UC15': ('done', 'Alerts page + processing'),
    'UC16': ('done', 'Export PDF/Excel'),
    'UC17': ('done', 'Báo cáo chi tiết lô'),
    'UC18': ('done', 'Compliance + TrustCheck'),
    'UC19': ('done', 'GTIN check API'),
    'UC20': ('done', 'Admin user management (suspend/delete)'),
}

# UC mới đã build (chưa có trong tài liệu gốc)
UC_NEW = [
    ('UC21', 'AI Photo Vision Verify', 'done', 'Client TF.js MobileNet 1024-dim features + cosine similarity'),
    ('UC22', 'Cấu hình Webhook ERP', 'done', 'HMAC-SHA256 outbound webhook + test ping UI'),
    ('UC23', 'Cấu hình Vòng đời theo DN', 'done', 'Per-DN threshold + custom theo ngành'),
    ('UC24', 'Geocoding DN address', 'done', 'VietMap Search API + batch script + cache lat/lng'),
    ('UC25', 'Time slider heatmap animation', 'done', '24h/7d/30d buckets + 4 speeds + auto-play'),
    ('UC26', 'Investigation Console', 'missing', 'Cần build: timeline vụ việc + assign cán bộ + evidence'),
    ('UC27', 'Side-by-side Image Comparison', 'missing', 'Cần build: 2 ảnh + AI diff % per component'),
    ('UC28', 'Reward System cho NTD', 'missing', 'Cần build: gamification + leaderboard + voucher'),
    ('UC29', 'Subscription Payment', 'missing', 'Cần build: 4 gói Basic/Pro/Enterprise/Government'),
    ('UC30', 'Unmask PII flow', 'missing', 'Cần build: yêu cầu cơ quan + approve pháp chế + decrypt'),
]

for uc_code, (status, detail) in UC_STATUS.items():
    idx, p = find_para(uc_code + ' ', 0)
    if not p:
        idx, p = find_para(uc_code + ' –', 0)
    if p:
        icon = {'done': '✅', 'partial': '🟡', 'missing': '❌'}.get(status, '•')
        color = {'done': GREEN, 'partial': AMBER, 'missing': RED}.get(status, SLATE)
        # Insert tiny status line right after UC heading
        insert_para_after(p, f'   {icon} {detail}', color=color, italic=True, size=9)
        print(f'  ✓ {uc_code}: {detail[:40]}')

print('=== STEP 4: Add NEW SECTIONS at end (from 5 detail docs) ===')
# Append at end of document

def add_h1(text, color=GOLD):
    """Add big heading manually (doc không có built-in Heading 1 style)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = color
    return p

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = GOLD
    return p

def add_h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    return p

def add_p(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def add_bullet(text):
    """Add bullet without relying on built-in style"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run('• ' + text)
    r.font.size = Pt(11)
    return p

def _set_cell_border(cell):
    """Add basic borders to a table cell"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '888888')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_table_simple(headers, rows):
    """Table without relying on built-in styles"""
    t = doc.add_table(rows=1, cols=len(headers))
    # Try built-in style, fallback to manual borders
    try:
        t.style = doc.styles['Table Grid']
    except KeyError:
        pass

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_border(hdr[i])
        # Shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '0B1623')
        hdr[i]._tc.get_or_add_tcPr().append(shading)

    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = ''
            p = row[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            _set_cell_border(row[i])
            # Color icons
            if str(val).startswith('✅'):
                run.font.color.rgb = GREEN
            elif str(val).startswith('❌') or str(val).startswith('🔴'):
                run.font.color.rgb = RED
            elif str(val).startswith('🟡') or str(val).startswith('🟠'):
                run.font.color.rgb = AMBER
    return t

doc.add_page_break()
add_h1('PHẦN BỔ SUNG TỪ 5 TÀI LIỆU CHI TIẾT MODULE (cập nhật 30/05/2026)', GOLD)
add_p(
    'Phần này tổng hợp các chi tiết hóa kiến trúc từ 5 tài liệu module gửi sau '
    '(file 1: Enterprise, file 2: Consumer, file 3: AI Engine, file 4: BI, file 5: Unified Portal). '
    'Nội dung dưới đây bổ sung những khái niệm CHƯA CÓ trong tài liệu gốc V2.0.',
    italic=True, color=SLATE
)

# ─── XI. UNIFIED SUPER PORTAL ─────────────────────────────────────
add_h1('XI. UNIFIED SUPER PORTAL (TỪ FILE 5)', GOLD)
add_p('🔴 CHƯA CÓ — Cần build landing page tập trung 4 nút lớn',
    bold=True, color=RED)

add_h2('11.1 Triết lý — Một nền tảng, nhiều vai trò')
add_p('Landing page tập trung với 4 nút lớn cho 4 module, giảm phân mảnh ứng dụng, '
    'tăng tốc truy cập, phù hợp định hướng IOC (Intelligent Operations Center) cấp quốc gia.')

add_h2('11.2 Layout chuẩn')
add_p('''
┌───────────────────────────────────────────────────┐
│              VNTRUST NATIONAL PLATFORM            │
│                                                   │
│  [1] DOANH NGHIỆP        [2] NGƯỜI TIÊU DÙNG     │
│      Enterprise Portal       Consumer App         │
│                                                   │
│  [3] AI KIỂM TRA         [4] PHÂN TÍCH/TRA CỨU   │
│      Fraud Detection         BI Command Center    │
│                                                   │
└───────────────────────────────────────────────────┘
''', size=10)

add_h2('11.3 Component bắt buộc')
add_table_simple(
    ['Khu vực', 'Nội dung'],
    [
        ('Header', 'Logo + Global Search + Notification + AI status + Role + Language'),
        ('Hero Center', '4 nút lớn vào 4 module (icon + tên + mô tả ngắn)'),
        ('Notification banner', 'Cảnh báo realtime (cảnh báo mới, QR clone, brand giả)'),
        ('Stats strip', 'KPI: DN verified, scans, fake alerts, active cases'),
        ('Footer', 'API docs link + hỗ trợ + chính sách'),
    ]
)

add_h2('11.4 ❌ 4 gói Subscription (HOÀN TOÀN MỚI — chưa có hệ thống payment)')
add_table_simple(
    ['Gói', 'Đối tượng', 'Quota QR/tháng', 'Tính năng', 'Phí'],
    [
        ('Basic', 'NSX nhỏ', '1,000', 'Hồ sơ + scan + lô', 'Free/Trial'),
        ('Pro', 'NSX trung', '10,000', 'Basic + API + Webhook', 'Paid'),
        ('Enterprise', 'NSX lớn', '100,000', 'Pro + AI Dashboard + IOC', 'Paid+'),
        ('Government', 'DN nhà nước', 'Unlimited', 'All + ưu tiên + audit', 'Custom'),
    ]
)

add_h2('11.5 PWA + Mobile-first')
add_bullet('Web App: React/Next.js + Tailwind (đã có ✅)')
add_bullet('PWA: cài đặt như app, hoạt động offline cơ bản (❌ cần manifest.json)')
add_bullet('Mobile native (giai đoạn 3): Flutter cho iOS/Android (❌)')
add_bullet('Edge AI: TensorFlow Lite chạy local cho QR verify offline (❌)')

# ─── XII. MODULE 1 BỔ SUNG ─────────────────────────────────────
add_h1('XII. MODULE 1 — ENTERPRISE PORTAL (BỔ SUNG TỪ FILE 1)', GOLD)

add_h2('12.1 ❌ Module Nhà máy & dây chuyền (chưa có model)')
add_p('Tài liệu file 1 yêu cầu lưu trữ chi tiết về nhà máy:',
    italic=True, color=AMBER)
add_table_simple(
    ['Nhóm', 'Dữ liệu cần có'],
    [
        ('Nhà máy', 'Địa chỉ + Tọa độ GPS + Hình ảnh + Video'),
        ('Dây chuyền', 'Loại + Công suất + Thiết bị'),
        ('Quy trình', 'SOP + GMP/HACCP file scan'),
        ('Đạo đức ESG/CSR', 'ESG report + Sustainability score + Carbon footprint'),
    ]
)
add_p('→ Cần thêm Prisma model: NhaMay, ESGReport', italic=True, color=RED)

add_h2('12.2 ❌ Module Thương hiệu / Logo riêng (chưa có)')
add_bullet('Tên thương hiệu + tên nhãn hiệu (khác nhau)')
add_bullet('Logo file (PNG/SVG) + mã màu hex chính thức + font chữ brand')
add_bullet('Số đăng ký nhãn hiệu (Cục SHTT)')
add_bullet('Phạm vi bảo hộ + ngày hết hạn')
add_p('→ Cần thêm Prisma model: ThuongHieu', italic=True, color=RED)

add_h2('12.3 ✅ Workflow phê duyệt 9 trạng thái (đã có 4 trạng thái cơ bản)')
add_table_simple(
    ['Trạng thái', 'Ý nghĩa', 'Status'],
    [
        ('Draft', 'Đang soạn, chưa gửi', '🟡 Có nhưng inline'),
        ('Submitted', 'Đã gửi chờ tiếp nhận', '✅'),
        ('Pending Review', 'Bộ phận kiểm tra đang xét', '✅ (pending)'),
        ('Need More Info', 'Yêu cầu bổ sung', '❌ Cần thêm'),
        ('Field Inspection', 'Kiểm tra thực tế nhà máy', '❌ Cần thêm'),
        ('Approved', 'Đã duyệt', '✅ (verified)'),
        ('Rejected', 'Từ chối có lý do', '✅ (suspended)'),
        ('Suspended', 'Tạm khóa vi phạm', '✅'),
        ('Revoked', 'Thu hồi vĩnh viễn', '🟡 (= delete)'),
    ]
)

# ─── XIII. MODULE 2 BỔ SUNG ─────────────────────────────────────
add_h1('XIII. MODULE 2 — CONSUMER CROWDSOURCING (BỔ SUNG TỪ FILE 2)', GOLD)

add_h2('13.1 ❌ Wizard 4 bước upload (chưa có UI wizard)')
add_p('Tài liệu yêu cầu UX wizard step-by-step thay vì single form:',
    italic=True, color=AMBER)
add_p('Bước 1 — Chọn nơi mua', bold=True)
add_bullet('Chợ truyền thống / Siêu thị / Tạp hóa / TMĐT / Livestream / Không rõ')
add_p('Bước 2 — Tình trạng sản phẩm', bold=True)
add_bullet('Nguyên seal / Mở hộp / Hư hỏng / Có dấu hiệu bất thường / Không có tem / Giá bất thường')
add_p('Bước 3 — Upload ảnh BẮT BUỘC', bold=True)
add_bullet('Mặt trước + Mặt sau + Tem chống giả + NSX/HSD + Barcode/QR')
add_p('Bước 4 — Xác nhận & gửi', bold=True)
add_bullet('GPS tự động + Device fingerprint + Hệ thống tạo mã ticket #CASE-YYYY-NNNNNN')

add_h2('13.2 ✅ Risk Scoring với weights cụ thể (đã có scoring nhưng weights khác)')
add_table_simple(
    ['Điều kiện', 'Điểm rủi ro +', 'Mức cảnh báo'],
    [
        ('QR không tồn tại', '+80', '🔴 Đỏ'),
        ('HSD sai định dạng', '+40', '🟡 Vàng'),
        ('Bao bì khác mẫu >50%', '+60', '🟠 Cam'),
        ('GPS bất thường (>500km vùng phân phối)', '+30', '🟡 Vàng'),
        ('Upload hàng loạt cùng device', '+50', '🟠 Cam'),
        ('Image duplicate đã upload', '+30', '🟡 Spam'),
    ]
)

add_h2('13.3 ❌ Reward System (chưa có)')
add_p('Gamification cho NTD báo cáo đúng:', italic=True)
add_bullet('Điểm thưởng cho mỗi báo cáo verified')
add_bullet('Leaderboard top reporter theo tháng/năm')
add_bullet('Đổi điểm → voucher / quà tặng từ DN tài trợ')
add_bullet('Badge: Sentinel, Inspector, Master Reporter')
add_p('→ Cần thêm Prisma models: UserScore, Reward', italic=True, color=RED)

add_h2('13.4 ❌ Livestream Detection (chưa có)')
add_bullet('AI giám sát stream bán hàng TikTok/Facebook/YouTube')
add_bullet('Phát hiện brand bị giả mạo trong livestream')
add_bullet('Tự động ghi nhận: shop name + product mention + timestamp')

add_h2('13.5 ❌ Social Media Monitoring (chưa có)')
add_bullet('Quét Facebook/Zalo/TikTok cho mention thương hiệu')
add_bullet('Sentiment analysis về DN/sản phẩm')
add_bullet('Phát hiện shop bán hàng giả qua post/comment')

add_h2('13.6 ❌ QR Dynamic (chưa có)')
add_p('QR token có TTL ngắn, refresh theo thời gian → chống clone',
    italic=True)

# ─── XIV. MODULE 3 BỔ SUNG ─────────────────────────────────────
add_h1('XIV. MODULE 3 — AI FRAUD DETECTION ENGINE (BỔ SUNG TỪ FILE 3)', GOLD)

add_h2('14.1 🟡 Triết lý kiểm tra 3 tầng (đã có partial)')
add_table_simple(
    ['Tầng', 'Loại kiểm tra', 'Thời gian', 'Status'],
    [
        ('1', 'So sánh dữ liệu cơ bản (text, mã, ngày)', '<100ms', '✅'),
        ('2', 'Logic nghiệp vụ (giá, GPS, batch, quota)', '<1s', '🟡 partial'),
        ('3', 'AI Vision & Similarity (logo, bao bì)', '2-10s', '🟡 client only'),
    ]
)

add_h2('14.2 ❌ Hệ thống 6 màu severity (đã có 4 màu)')
add_table_simple(
    ['Màu', 'Mức', 'Điểm', 'Ý nghĩa', 'Status'],
    [
        ('🟢 Xanh lá', 'PASS', '0-20', 'Đạt', '✅'),
        ('🟡 Vàng', 'MONITOR', '21-40', 'Theo dõi', '✅'),
        ('🟠 Cam', 'SUSPECT', '41-60', 'Nghi vấn', '❌ Chưa có'),
        ('🔴 Đỏ', 'RISK', '61-80', 'Rủi ro cao', '✅'),
        ('⚫ Đen', 'FRAUD', '81-100', 'Giả mạo xác định', '❌ Chưa có'),
        ('🔵 Xanh dương', 'UNKNOWN', '—', 'Chưa đủ data', '❌ Chưa có'),
    ]
)

add_h2('14.3 ❌ Checklist UI (chưa có)')
add_p('Hiển thị từng tiêu chí check với màu sắc + kết luận tổng:',
    italic=True)
add_p('''
RISK ASSESSMENT CHECKLIST                Risk Score: 87/100  [⚫ FRAUD]
─────────────────────────────────────────────────────────────────
[🟢] Mã doanh nghiệp hợp lệ
[🟢] Chứng nhận ISO còn hiệu lực
[🟡] Giá bán thấp bất thường (-45%)
[🟠] Bao bì có 72% tương đồng với mẫu chuẩn
[🔴] QR bị quét tại 5 tỉnh trong 1 giờ
[🔴] Logo khác mẫu chuẩn 38%
[⚫] Barcode không tồn tại trong GS1

KẾT LUẬN: NGHI HÀNG GIẢ — Chuyển QLTT điều tra
''', size=9)

add_h2('14.4 🟡 6 thuật toán similarity (đã có cosine, thiếu 5)')
add_table_simple(
    ['Thuật toán', 'Mục tiêu', 'Status'],
    [
        ('Levenshtein', 'Sai khác ký tự (tên DN, brand)', '❌'),
        ('Cosine Similarity', 'Tương đồng văn bản', '✅'),
        ('Fuzzy Matching', 'Gần giống có lỗi typo', '❌'),
        ('Soundex', 'Phát âm tương tự', '❌'),
        ('NLP Embedding', 'Ngữ nghĩa (BERT)', '❌'),
        ('Siamese CNN', 'Image similarity sâu', '❌'),
    ]
)

add_h2('14.5 🔴 AI Vision Models — chỉ có MobileNet client')
add_table_simple(
    ['Model', 'Chức năng', 'Status hiện tại'],
    [
        ('CNN (MobileNet)', 'Feature extraction', '✅ Có client-side'),
        ('Siamese Network', 'Image similarity 2-vector', '❌ Chưa có'),
        ('YOLOv8', 'Object/logo detection', '❌ Chưa có'),
        ('OCR (PaddleOCR)', 'Đọc nhãn tự động', '🟡 Stub tesseract.js'),
        ('Segmentation (SAM)', 'Tách vùng bao bì', '❌ Chưa có'),
        ('GAN Detection', 'Phát hiện deepfake packaging', '❌ Chưa có'),
    ]
)

add_h2('14.6 ❌ Bao bì comparison weights chuẩn')
add_table_simple(
    ['Thành phần', 'Trọng số', 'Ghi chú'],
    [
        ('Logo', '25%', 'Quan trọng nhất'),
        ('Tem chống giả', '20%', 'Đặc trưng nhất'),
        ('Bố cục layout', '15%', 'Khó copy chính xác'),
        ('Font chữ', '15%', 'Brand identity'),
        ('Mã QR', '15%', 'Phải khớp'),
        ('Màu sắc', '10%', 'Có sai số do ánh sáng'),
    ]
)

add_h2('14.7 🟡 Hybrid Review Workflow (đã có implicit)')
add_p('''
AI phân tích
      ↓
Risk Score
      ↓
< 40   →  Tự động thông qua (auto pass) ✅
40-70  →  Chuyển kiểm tra nhân công ❌ (chưa có UI)
> 70   →  Cảnh báo ĐỎ + Điều tra thực địa ❌ (chưa có Investigation Console)
''', size=10)

add_h2('14.8 ❌ Investigation Console (chưa có)')
add_bullet('Timeline vụ việc với từng action (upload → AI alert → assign QLTT → kết luận)')
add_bullet('Side-by-side image comparison (ảnh nghi vấn vs ảnh gốc) + diff %')
add_bullet('Notes, evidence attachments, status changes')
add_bullet('Assign cán bộ điều tra + due date + escalation')
add_bullet('Export báo cáo vụ việc PDF')
add_p('→ Cần Prisma models: Investigation/Case + InvestigationNote + Evidence',
    italic=True, color=RED)

# ─── XV. MODULE 4 BỔ SUNG ─────────────────────────────────────
add_h1('XV. MODULE 4 — BI/ANALYTICS (BỔ SUNG TỪ FILE 4)', GOLD)

add_h2('15.1 ✅ 12 mục menu Command Center (đã có 10/12)')
items_status = [
    ('1. Dashboard tổng quan', '✅'),
    ('2. Tra cứu nhanh', '🟡 basic'),
    ('3. Phân tích hàng giả', '✅'),
    ('4. Phân tích doanh nghiệp', '✅'),
    ('5. Heatmap GPS', '✅ (VietMap)'),
    ('6. Theo dõi thương hiệu', '❌'),
    ('7. AI Risk Analytics', '🟡'),
    ('8. Điều tra & vụ việc', '❌'),
    ('9. Báo cáo thống kê', '✅'),
    ('10. Cảnh báo thời gian thực', '✅'),
    ('11. Xuất dữ liệu / API', '✅'),
    ('12. Quản trị hệ thống', '✅ (system-config)'),
]
add_table_simple(['Mục', 'Status'], items_status)

add_h2('15.2 🟡 8 loại biểu đồ (đã có 4/8)')
add_table_simple(
    ['Loại', 'Status', 'Use case'],
    [
        ('Bar Chart', '✅', 'Top 10 SP bị giả'),
        ('Line Chart', '✅', 'Số scan/ngày'),
        ('Pie Chart', '✅', 'Cơ cấu vi phạm'),
        ('Heatmap GIS', '✅', 'Bản đồ điểm nóng'),
        ('Sankey', '❌', 'Dòng phân phối lô hàng'),
        ('Network Graph', '❌', 'Mạng lưới shop liên kết'),
        ('Funnel', '❌', 'Quét → báo cáo → xử lý'),
        ('Treemap', '❌', 'Cơ cấu ngành/sub-ngành'),
    ]
)

add_h2('15.3 ✨ Time slider animation (BONUS — vượt yêu cầu tài liệu)')
add_bullet('Slider chia thành 24h / 7d / 30d buckets')
add_bullet('Auto-play 4 tốc độ (×½, ×1, ×2, ×3)')
add_bullet('Heatmap evolve theo thời gian → phát hiện xu hướng lan rộng')

add_h2('15.4 ✅ GIS Heatmap với 4 layers (đáp ứng + nhiều hơn)')
add_table_simple(
    ['Layer', 'Mô tả', 'Status'],
    [
        ('Scan layer', 'Mật độ lượt quét', '✅'),
        ('Fake layer', 'Mật độ hàng giả', '✅'),
        ('Alert layer', 'Cảnh báo mở', '✅'),
        ('DN layer', 'Phân bố doanh nghiệp', '✅ (geocoded)'),
    ]
)

add_h2('15.5 ❌ Export Power BI / Tableau (chưa có connector)')

# ─── XVI. APPENDIX TỔNG QUAN ─────────────────────────────────────
add_h1('XVI. PHỤ LỤC — TỔNG QUAN COMPLIANCE & TECH STACK', GOLD)

add_h2('16.1 Mức tuân thủ theo module')
add_table_simple(
    ['Module / Phần', 'Compliance', 'Trạng thái'],
    [
        ('Module 1 — Enterprise Portal', '85%', '🟢 Tốt, thiếu Nhà máy/Logo/ESG'),
        ('Module 2 — Consumer Crowdsourcing', '75%', '🟡 Thiếu Wizard, Reward, Livestream'),
        ('Module 3 — AI Fraud Detection', '60%', '🟠 Thiếu Investigation, Side-by-side, YOLOv8'),
        ('Module 4 — BI/Analytics', '80%', '🟢 Có Heatmap + Time slider, thiếu Sankey/Network'),
        ('Module 5 — Unified Portal', '40%', '🔴 Thiếu Landing 4 nút + Payment'),
        ('Phân quyền + KYC + UC03', '95%', '🟢 Đầy đủ'),
        ('Bảo mật + Vault', '85%', '🟢 Có vault, thiếu Unmask UI'),
        ('Cảnh báo vòng đời + Hậu kiểm', '95%', '🟢 Đầy đủ'),
        ('TỔNG THỂ', '~83%', '🟢 Production-ready'),
    ]
)

add_h2('16.2 Use Cases mới CẦN BUILD (UC26-UC30)')
add_table_simple(
    ['Code', 'Tên', 'Effort'],
    [
        ('UC26', 'Investigation Console (timeline + assign + evidence)', '3 ngày'),
        ('UC27', 'Side-by-side Image Comparison + diff %', '2 ngày'),
        ('UC28', 'Reward System cho NTD (gamification + leaderboard)', '3 ngày'),
        ('UC29', 'Subscription Payment (Basic/Pro/Enterprise/Gov)', '5-7 ngày'),
        ('UC30', 'Unmask PII flow cho cơ quan chức năng', '2 ngày'),
        ('UC31 (mới)', 'Unified Super Portal landing 4 nút', '1-2 ngày'),
    ]
)

add_h2('16.3 Tech Stack hiện tại vs Recommended')
add_table_simple(
    ['Component', 'Hiện tại', 'Recommended từ docs'],
    [
        ('Frontend', 'Next.js 16 + React 19 + Tailwind ✅', 'React/NextJS ✅'),
        ('Backend API', 'Next.js API Routes', 'NestJS ❌'),
        ('Database', 'SQLite (dev), PostgreSQL (prod) ✅', 'PostgreSQL ✅'),
        ('Auth', 'JWT HS256 + bcrypt ✅', 'OAuth2 + SSO 🟡'),
        ('Map', 'MapLibre + CartoDB + Esri ✅', 'Mapbox ≈'),
        ('AI Client', 'TF.js MobileNet ✅', 'TensorFlow ✅'),
        ('AI Server', 'Không có ❌', 'YOLOv8 + Siamese + GAN ❌'),
        ('OCR', 'Tesseract.js stub 🟡', 'PaddleOCR ❌'),
        ('Workflow', 'Custom Prisma logic', 'Camunda BPMN ❌'),
        ('Search', 'Prisma queries 🟡', 'Elasticsearch + Vector ❌'),
        ('Stream', 'Không có ❌', 'Kafka ❌'),
        ('Log analytics', 'NhatKy table 🟡', 'ClickHouse ❌'),
        ('BI', 'Custom dashboards ✅', 'Apache Superset ❌'),
        ('Email', 'Nodemailer + Gmail ✅', 'SendGrid/SES ≈'),
    ]
)

add_h2('16.4 21 Prisma models hiện có')
models_now = [
    'DoanhNghiep', 'SanPham', 'LoHang', 'KhoHang', 'MaDinhDanh', 'LuotQuet',
    'CanhBao', 'ChungNhan', 'NguoiDung', 'LoiMoiNhanVien', 'KetQuaHauKiem',
    'CauHinhHeThong', 'YeuCauTuanThuVSIC', 'IdentityVault', 'ReportVault',
    'AnonymousSession', 'WebhookErp', 'TieuChuanKiemNghiem', 'NhatKy', 'ThongBao',
    'DonChuyenHang',
]
for m in models_now:
    add_bullet(f'✅ {m}')

add_h2('16.5 ❌ 9 Models cần BỔ SUNG (từ 5 file mới)')
add_table_simple(
    ['Model mới', 'Mục đích', 'Cần cho'],
    [
        ('NhaMay', 'Nhà máy + dây chuyền + SOP', 'Module 1'),
        ('ThuongHieu', 'Logo + nhãn hiệu + bảo hộ SHTT', 'Module 1'),
        ('ESGReport', 'Báo cáo ESG/CSR của DN', 'Module 1'),
        ('Investigation', 'Case timeline + evidence', 'Module 3+4'),
        ('UserScore', 'Gamification điểm thưởng NTD', 'Module 2'),
        ('Reward', 'Voucher/quà tặng', 'Module 2'),
        ('Subscription', 'Gói phí Basic/Pro/Ent/Gov', 'Module 5'),
        ('Payment', 'Lịch sử thanh toán + invoice', 'Module 5'),
        ('SocialMention', 'Monitoring MXH', 'Module 2'),
        ('UnmaskRequest', 'Yêu cầu giải mã PII', 'Module 3'),
    ]
)

add_h2('16.6 Roadmap 3 giai đoạn')
add_h3('Giai đoạn 1 — MVP (Đã hoàn thành ~85%)')
add_bullet('✅ Enterprise Portal cơ bản + KYC')
add_bullet('✅ Consumer scan + report')
add_bullet('✅ AI verify image basic')
add_bullet('✅ Heatmap 4 layers + Time slider')
add_bullet('✅ Webhook ERP + Lifecycle config + System config')

add_h3('Giai đoạn 2 — Nâng cao (Q3/2026)')
add_bullet('🟡 Landing Unified Portal 4 nút (UC31)')
add_bullet('🟡 Investigation Console (UC26)')
add_bullet('🟡 Side-by-side image comparison (UC27)')
add_bullet('🟡 6-color checklist UI')
add_bullet('🟡 Subscription Payment (UC29)')
add_bullet('🟡 Wizard Consumer + Reward (UC28)')
add_bullet('🟡 Unmask PII flow (UC30)')
add_bullet('🟡 Server-side AI (YOLOv8 + PaddleOCR + Siamese)')
add_bullet('🟡 Mobile PWA installable')

add_h3('Giai đoạn 3 — Mở rộng (Q1/2027)')
add_bullet('🔵 Mobile native Flutter')
add_bullet('🔵 Livestream Detection')
add_bullet('🔵 Social Media Monitoring')
add_bullet('🔵 Blockchain traceability')
add_bullet('🔵 Edge AI offline')
add_bullet('🔵 National IOC dashboard')
add_bullet('🔵 Integration QLTT/Hải quan/ATTP nationwide')

# Footer note
doc.add_paragraph()
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot_run = foot.add_run('— HẾT TÀI LIỆU (UPDATED 30/05/2026 — 44 tasks hoàn thành) —')
foot_run.italic = True
foot_run.font.color.rgb = GOLD
foot_run.font.size = Pt(11)

# Save
doc.save(FILE)
print(f'\n✅ Saved (in-place edit): {FILE}')
print(f'   Size: {os.path.getsize(FILE)/1024:.1f} KB')
print(f'   Paragraphs: {len(doc.paragraphs)}')
print(f'   Tables: {len(doc.tables)}')
