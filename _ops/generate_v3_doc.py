"""
Generate TAI_LIEU_NGHIEP_VU_VNTRUST_V3.docx
Hợp nhất: tài liệu chính + 5 file chi tiết module + status hệ thống built + gap analysis
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'C:\xampp\htdocs\Web-chong-hang-gia-main'
OUT = os.path.join(ROOT, 'TAI_LIEU_NGHIEP_VU_VNTRUST_V3.docx')

# Color palette
GOLD = RGBColor(0xC8, 0xA5, 0x57)
INK = RGBColor(0x0B, 0x16, 0x23)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED = RGBColor(0xEF, 0x44, 0x44)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
SLATE = RGBColor(0x47, 0x55, 0x69)
BLUE = RGBColor(0x3B, 0x82, 0xF6)

doc = Document()

# ─── Page setup ─────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ─── Helper functions ────────────────────────────────────────────────

def add_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    return p

def add_number(text):
    return doc.add_paragraph(text, style='List Number')

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        # Shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '0B1623')
        hdr[i]._tc.get_or_add_tcPr().append(shading)

    # Rows
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            cell = row[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            # Color status icons
            if val in ('✅', '✓'):
                run.font.color.rgb = GREEN
                run.bold = True
            elif val in ('❌', '✗', '🔴'):
                run.font.color.rgb = RED
                run.bold = True
            elif val in ('🟡', '🟠'):
                run.font.color.rgb = AMBER
                run.bold = True

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    return t

def add_page_break():
    doc.add_page_break()

def add_divider(text=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50 + (f' {text} ' if text else '') + '─' * 50)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)

# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('TÀI LIỆU NGHIỆP VỤ HỢP NHẤT')
title_run.font.size = Pt(28)
title_run.bold = True
title_run.font.color.rgb = INK

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run('HỆ THỐNG VNTRUST')
sub_run.font.size = Pt(36)
sub_run.bold = True
sub_run.font.color.rgb = GOLD

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_run = ver.add_run('Phiên bản 3.0 — Hợp nhất 6 tài liệu nguồn')
ver_run.font.size = Pt(14)
ver_run.italic = True
ver_run.font.color.rgb = SLATE

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc.add_run('Chống hàng giả & Xác thực nguồn gốc sản phẩm Việt Nam')
desc_run.font.size = Pt(13)
desc_run.font.color.rgb = SLATE

doc.add_paragraph()
doc.add_paragraph()

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Ngày tổng hợp', '30/05/2026'),
    ('Tài liệu nguồn', '6 file: tài liệu chính + 5 file chi tiết module'),
    ('Tình trạng tuân thủ', '~83% (96% framework, 70% chi tiết module)'),
    ('Bảo mật', 'Nội bộ — Không phát tán'),
]
for i, (k, v) in enumerate(meta_data):
    meta_table.rows[i].cells[0].text = k
    meta_table.rows[i].cells[1].text = v
    for cell in meta_table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
    meta_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# MỤC LỤC
# ═══════════════════════════════════════════════════════════════════════
add_heading('MỤC LỤC', 0, GOLD)
toc_items = [
    'I. TỔNG QUAN HỆ THỐNG',
    'II. KIẾN TRÚC UNIFIED SUPER PORTAL (Mới)',
    'III. MODULE 1 — ENTERPRISE PORTAL',
    'IV. MODULE 2 — CONSUMER CROWDSOURCING',
    'V. MODULE 3 — AI FRAUD DETECTION ENGINE',
    'VI. MODULE 4 — BI/ANALYTICS COMMAND CENTER',
    'VII. PHÂN QUYỀN, KYC & SUBSCRIPTION',
    'VIII. BẢO MẬT, VAULT & PHÁP LÝ',
    'IX. USE CASES (UC01–UC30)',
    'X. CẢNH BÁO VÒNG ĐỜI & HẬU KIỂM',
    'XI. TRẠNG THÁI HIỆN TẠI & GAP ANALYSIS',
    'XII. ROADMAP 3 GIAI ĐOẠN',
]
for i, item in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.add_run(f'{i}. {item}').font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# I. TỔNG QUAN HỆ THỐNG
# ═══════════════════════════════════════════════════════════════════════
add_heading('I. TỔNG QUAN HỆ THỐNG', 1, GOLD)

add_heading('1.1 Mục tiêu chiến lược', 2)
doc.add_paragraph(
    'VNTrust là nền tảng "Unified Anti-Counterfeit Intelligence Platform" — hệ thống '
    'điều hành tập trung cho việc xác thực doanh nghiệp, sản phẩm, chống hàng giả, '
    'giám sát thị trường, và truy xuất nguồn gốc theo chuẩn Chính phủ điện tử.'
)
add_bullet('Quản lý doanh nghiệp sản xuất / nhập khẩu với workflow phê duyệt KYC')
add_bullet('Phát hiện hàng giả qua AI Vision + crowdsourcing người tiêu dùng')
add_bullet('Bản đồ điểm nóng GIS với heatmap thời gian thực')
add_bullet('Tích hợp API với cơ quan QLTT, Hải quan, ATTP')
add_bullet('Tuân thủ Luật Bảo vệ DLCN 2025, ATTT ISO 27001, OAuth2 + PKI')

add_heading('1.2 Đối tượng người dùng (6 nhóm)', 2)
add_table(
    ['Vai trò', 'Quyền chính'],
    [
        ('Doanh nghiệp (NSX/NNK)', 'Khai báo hồ sơ + cập nhật SP/lô/QR'),
        ('Bộ phận kiểm duyệt', 'Phê duyệt KYC + workflow chứng từ'),
        ('Quản trị hệ thống (Admin)', 'Toàn quyền + cấu hình + audit'),
        ('Cơ quan quản lý (QLTT/Hải quan/ATTP)', 'Read-only + unmask PII có lệnh'),
        ('Đối tác API', 'Webhook ERP + REST API tích hợp'),
        ('Người tiêu dùng', 'Quét QR + báo cáo nghi vấn + ẩn danh'),
    ],
    col_widths=[6, 11]
)

add_heading('1.3 Kiến trúc 4 module + 1 portal', 2)
doc.add_paragraph(
    'Hệ thống chia thành 4 module độc lập + 1 portal hợp nhất (theo tài liệu chi tiết file 5):'
)
add_bullet('Module 1: Enterprise Portal — Hồ sơ DN, chứng nhận, sản phẩm, workflow phê duyệt')
add_bullet('Module 2: Consumer Crowdsourcing — Quét QR, chụp ảnh, báo cáo, GPS')
add_bullet('Module 3: AI Fraud Detection Engine — Rule + Vision + Similarity + Scoring')
add_bullet('Module 4: BI/Analytics Command Center — Search, Heatmap, Reports, IOC')
add_bullet('Module 5 (Portal): Unified landing với 4 nút lớn dẫn vào 4 module')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# II. UNIFIED SUPER PORTAL
# ═══════════════════════════════════════════════════════════════════════
add_heading('II. KIẾN TRÚC UNIFIED SUPER PORTAL', 1, GOLD)

add_heading('2.1 Triết lý — Một nền tảng, nhiều vai trò', 2)
doc.add_paragraph(
    'Landing page tập trung với 4 nút lớn cho 4 module, giảm phân mảnh ứng dụng, '
    'tăng tốc truy cập, phù hợp định hướng IOC (Intelligent Operations Center).'
)

add_heading('2.2 Layout landing page', 2)
p = doc.add_paragraph()
run = p.add_run('''
┌───────────────────────────────────────────────────┐
│              VNTRUST NATIONAL PLATFORM            │
│                                                   │
│  [1] DOANH NGHIỆP        [2] NGƯỜI TIÊU DÙNG     │
│      Enterprise Portal       Consumer App         │
│                                                   │
│  [3] AI KIỂM TRA         [4] PHÂN TÍCH/TRA CỨU   │
│      Fraud Detection         BI Command Center    │
│                                                   │
└───────────────────────────────────────────────────┘
''')
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_heading('2.3 Component bắt buộc', 2)
add_table(
    ['Khu vực', 'Nội dung'],
    [
        ('Header', 'Logo + Global Search + Notification bell + AI status + Role + Language'),
        ('Hero Center', '4 nút lớn vào 4 module (icon + tên + mô tả ngắn)'),
        ('Notification banner', 'Cảnh báo realtime (cảnh báo mới, QR clone, brand bị giả)'),
        ('Stats strip', 'KPI: DN verified, scans, fake alerts, active cases'),
        ('Footer', 'API docs link + hỗ trợ + chính sách'),
    ],
    col_widths=[5, 12]
)

add_heading('2.4 4 gói subscription (NEW)', 2)
add_table(
    ['Gói', 'Đối tượng', 'Quyền lợi', 'Phí tham khảo'],
    [
        ('Basic', 'NSX nhỏ', 'Hồ sơ + 1000 QR/tháng', 'Free / Trial'),
        ('Pro', 'NSX trung', 'Hồ sơ + 10K QR + API + Webhook', 'Paid'),
        ('Enterprise', 'NSX lớn', 'AI Dashboard + IOC + multi-user', 'Paid+'),
        ('Government Verified', 'DN nhà nước', 'Xác thực cao + ưu tiên', 'Custom'),
    ],
    col_widths=[4, 4, 6, 3]
)

add_heading('2.5 PWA + Mobile-first', 2)
add_bullet('Web App: React/Next.js + Tailwind dark theme gold accent')
add_bullet('PWA: cài đặt như app, hoạt động offline cơ bản')
add_bullet('Mobile native (giai đoạn 3): Flutter cho iOS/Android')
add_bullet('Edge AI: TensorFlow Lite chạy local cho QR verify offline')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# III. MODULE 1 — ENTERPRISE
# ═══════════════════════════════════════════════════════════════════════
add_heading('III. MODULE 1 — ENTERPRISE PORTAL', 1, GOLD)

add_heading('3.1 Mục tiêu', 2)
doc.add_paragraph(
    'Quản lý đầy đủ hồ sơ doanh nghiệp sản xuất/nhập khẩu: từ pháp lý, nhà máy, dây chuyền, '
    'chứng nhận, sản phẩm, bao bì đến QR/API và workflow phê duyệt.'
)

add_heading('3.2 Sidebar menu DN (12 mục)', 2)
menu_items = [
    '1. Tổng quan (Dashboard)', '2. Hồ sơ doanh nghiệp', '3. Nhà máy & sản xuất',
    '4. Chứng nhận & tiêu chuẩn', '5. Giấy phép lưu hành', '6. Sản phẩm',
    '7. Bao bì & truy xuất', '8. Mã QR / API', '9. Hồ sơ đang xử lý',
    '10. Lịch sử kiểm tra', '11. Cảnh báo & vi phạm', '12. Tài khoản & phân quyền',
]
for m in menu_items:
    add_bullet(m)

add_heading('3.3 Form hồ sơ pháp lý — bắt buộc', 2)
add_table(
    ['Nhóm', 'Trường dữ liệu'],
    [
        ('Pháp lý', 'Tên công ty, MST, Giấy ĐKKD, Ngày cấp, Nơi cấp'),
        ('Liên hệ', 'Địa chỉ, GPS/Tọa độ, Email, Website, Hotline'),
        ('Đại diện', 'Họ tên, CCCD/Hộ chiếu, Chức vụ, Chữ ký số (PKI)'),
    ],
    col_widths=[4, 13]
)

add_heading('3.4 Module Nhà máy & dây chuyền (MỚI — chưa có trong system)', 2, RED)
add_table(
    ['Nhóm', 'Dữ liệu'],
    [
        ('Nhà máy', 'Địa chỉ, Tọa độ GPS, Hình ảnh, Video'),
        ('Dây chuyền', 'Loại, Công suất, Thiết bị'),
        ('Quy trình', 'SOP, GMP/HACCP'),
        ('Đạo đức ESG/CSR', 'ESG report, CSR cam kết, Sustainability score'),
    ],
    col_widths=[4, 13]
)

add_heading('3.5 Module Thương hiệu / Logo (MỚI — chưa có model riêng)', 2, RED)
add_bullet('Tên thương hiệu + tên nhãn hiệu')
add_bullet('Logo file (PNG/SVG) + mã màu (hex) + font chữ chính thức')
add_bullet('Số đăng ký nhãn hiệu (Cục SHTT)')
add_bullet('Phạm vi bảo hộ + ngày hết hạn')

add_heading('3.6 Workflow phê duyệt (9 trạng thái)', 2)
add_table(
    ['Trạng thái', 'Ý nghĩa'],
    [
        ('Draft', 'Đang soạn, chưa gửi'),
        ('Submitted', 'Đã gửi chờ tiếp nhận'),
        ('Pending Review', 'Bộ phận kiểm tra đang xét'),
        ('Need More Info', 'Yêu cầu bổ sung'),
        ('Field Inspection', 'Kiểm tra thực tế tại nhà máy'),
        ('Approved', 'Đã duyệt'),
        ('Rejected', 'Từ chối có lý do'),
        ('Suspended', 'Tạm khóa vi phạm'),
        ('Revoked', 'Thu hồi vĩnh viễn'),
    ],
    col_widths=[5, 12]
)

add_heading('3.7 6 API công khai cho đối tác', 2)
add_table(
    ['API', 'Chức năng'],
    [
        ('Company API', 'Tra cứu thông tin DN theo MST'),
        ('Product API', 'Truy xuất SP theo SKU/GTIN'),
        ('QR Verify API', 'Kiểm tra QR có hợp lệ không'),
        ('Certificate API', 'Xác thực chứng nhận theo số'),
        ('Alert API', 'Lấy cảnh báo theo DN/khu vực'),
        ('Audit API', 'Lấy nhật ký hoạt động'),
    ],
    col_widths=[5, 12]
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# IV. MODULE 2 — CONSUMER
# ═══════════════════════════════════════════════════════════════════════
add_heading('IV. MODULE 2 — CONSUMER CROWDSOURCING', 1, GOLD)

add_heading('4.1 Mục tiêu', 2)
doc.add_paragraph(
    'Biến mỗi NTD thành "thanh tra di động", thu thập dữ liệu thị trường real-time. '
    'Xây dựng bản đồ rủi ro hàng giả toàn quốc qua crowdsourcing intelligence.'
)

add_heading('4.2 6 hình thức đăng nhập', 2)
add_table(
    ['Hình thức', 'Mục tiêu'],
    [
        ('Username/Password', 'Tài khoản chuẩn dài hạn'),
        ('OTP điện thoại', 'Xác thực nhanh, 1 chạm'),
        ('Google / Apple OAuth', 'Social login'),
        ('CCCD/VNeID', 'Định danh quốc gia'),
        ('Guest', 'Ẩn danh tuyệt đối, không lưu danh tính'),
        ('Anonymous Protected', 'Mã hóa định danh, có thể unmask khi có lệnh'),
    ],
    col_widths=[5, 12]
)

add_heading('4.3 Wizard 4 bước upload (MỚI — chưa có wizard UI)', 2, RED)
add_number('Bước 1: Chọn nơi mua (Chợ / Siêu thị / Tạp hóa / TMĐT / Livestream / Không rõ)')
add_number('Bước 2: Tình trạng (Nguyên seal / Mở hộp / Hư hỏng / Không tem / Giá bất thường)')
add_number('Bước 3: Upload ảnh bắt buộc — Mặt trước, mặt sau, tem, NSX/HSD, QR/Barcode')
add_number('Bước 4: Mô tả + GPS tự động + xác nhận gửi')

add_heading('4.4 AI Risk Scoring với weights (chuẩn hóa)', 2)
add_table(
    ['Điều kiện', 'Điểm rủi ro +', 'Mức cảnh báo'],
    [
        ('QR không tồn tại', '+80', '🔴 Đỏ'),
        ('HSD sai định dạng', '+40', '🟡 Vàng'),
        ('Bao bì khác mẫu chuẩn >50%', '+60', '🟠 Cam'),
        ('GPS bất thường (Geo distance > 500km)', '+30', '🟡 Vàng'),
        ('Upload hàng loạt cùng device', '+50', '🟠 Cam'),
        ('Image duplicate (đã upload trước)', '+30', '🟡 Spam'),
    ],
    col_widths=[8, 4, 5]
)

add_heading('4.5 5 trạng thái kết quả', 2)
add_table(
    ['Trạng thái', 'Ý nghĩa', 'Hành động'],
    [
        ('VERIFIED', 'Chính hãng', 'Hiện đầy đủ info SP'),
        ('WARNING', 'Nghi vấn', 'Hỏi NTD xác nhận thêm'),
        ('EXPIRED', 'Hết hạn', 'Cảnh báo không tiêu thụ'),
        ('UNKNOWN', 'Chưa xác minh', 'Tạo ticket kiểm tra'),
        ('BLOCKED', 'Vi phạm xác định', 'Báo cơ quan QLTT'),
    ],
    col_widths=[4, 5, 8]
)

add_heading('4.6 Tính năng MỚI cần bổ sung', 2, RED)
add_bullet('🔴 Reward System — thưởng NTD báo cáo đúng (gamification, leaderboard)')
add_bullet('🔴 Livestream Detection — AI giám sát stream bán hàng TikTok/FB')
add_bullet('🔴 Social Media Monitoring — quét Facebook/Zalo cho mention thương hiệu')
add_bullet('🔴 QR Dynamic — QR thay đổi theo thời gian chống clone')
add_bullet('🔴 Edge AI offline — TensorFlow Lite verify QR không cần internet')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# V. MODULE 3 — AI ENGINE
# ═══════════════════════════════════════════════════════════════════════
add_heading('V. MODULE 3 — AI FRAUD DETECTION ENGINE', 1, GOLD)

add_heading('5.1 Triết lý kiểm tra 3 tầng', 2)
add_table(
    ['Tầng', 'Loại kiểm tra', 'Thời gian'],
    [
        ('1', 'So sánh dữ liệu cơ bản (text, mã, ngày)', 'Tức thì <100ms'),
        ('2', 'Logic nghiệp vụ (giá, GPS, batch, quota)', '< 1s'),
        ('3', 'AI Vision & Similarity (logo, bao bì, font)', '2-10s'),
    ],
    col_widths=[2, 10, 5]
)

add_heading('5.2 Hệ thống 6 màu severity (MỚI — chuẩn 5-6 màu)', 2, RED)
add_table(
    ['Màu', 'Mức', 'Điểm', 'Ý nghĩa'],
    [
        ('🟢 Xanh lá', 'PASS', '0-20', 'Đạt, không vấn đề'),
        ('🟡 Vàng', 'MONITOR', '21-40', 'Theo dõi'),
        ('🟠 Cam', 'SUSPECT', '41-60', 'Nghi vấn cần xem'),
        ('🔴 Đỏ', 'RISK', '61-80', 'Rủi ro cao'),
        ('⚫ Đen', 'FRAUD', '81-100', 'Giả mạo xác định'),
        ('🔵 Xanh dương', 'UNKNOWN', '—', 'Chưa đủ dữ liệu'),
    ],
    col_widths=[3, 3, 3, 8]
)

add_heading('5.3 Checklist UI (MỚI — chưa có)', 2, RED)
p = doc.add_paragraph()
run = p.add_run('''
RISK ASSESSMENT CHECKLIST                Risk Score: 87/100  [⚫ FRAUD]
─────────────────────────────────────────────────────────────────
[🟢] Mã doanh nghiệp hợp lệ
[🟢] Chứng nhận ISO còn hiệu lực
[🟡] Giá bán thấp bất thường (-45%)
[🟠] Bao bì có 72% tương đồng với mẫu chuẩn
[🔴] QR bị quét tại 5 tỉnh trong 1 giờ
[🔴] Logo khác mẫu chuẩn 38%
[⚫] Barcode không tồn tại trong GS1

KẾT LUẬN: NGHI HÀNG GIẢ — Chuyển QLTT điều tra
''')
run.font.name = 'Consolas'
run.font.size = Pt(9)

add_heading('5.4 6 thuật toán similarity', 2)
add_table(
    ['Thuật toán', 'Mục tiêu', 'Đối tượng'],
    [
        ('Levenshtein', 'Sai khác ký tự', 'Tên DN, brand'),
        ('Cosine Similarity', 'Tương đồng văn bản', 'Mô tả, slogan'),
        ('Fuzzy Matching', 'Gần giống có lỗi typo', 'SKU, mã'),
        ('Soundex', 'Phát âm tương tự', 'Tên Việt → English'),
        ('NLP Embedding', 'Ngữ nghĩa', 'Mô tả SP'),
        ('Siamese CNN', 'Image similarity sâu', 'Logo, bao bì'),
    ],
    col_widths=[4, 6, 7]
)

add_heading('5.5 AI Vision Models (MỚI — chỉ có MobileNet client)', 2, RED)
add_table(
    ['Model', 'Chức năng', 'Status'],
    [
        ('CNN (MobileNet)', 'Feature extraction', '✅ Có client-side'),
        ('Siamese Network', 'Image similarity 2-vector', '❌ Chưa có'),
        ('YOLOv8', 'Object/logo detection', '❌ Chưa có'),
        ('OCR (PaddleOCR)', 'Đọc nhãn tự động', '🟡 Stub tesseract'),
        ('Segmentation (SAM)', 'Tách vùng bao bì', '❌ Chưa có'),
        ('GAN Detection', 'Phát hiện deepfake packaging', '❌ Chưa có'),
    ],
    col_widths=[5, 7, 5]
)

add_heading('5.6 Bao bì comparison weights chuẩn', 2)
add_table(
    ['Thành phần', 'Trọng số', 'Ghi chú'],
    [
        ('Logo', '25%', 'Quan trọng nhất'),
        ('Tem chống giả', '20%', 'Đặc trưng nhất'),
        ('Bố cục layout', '15%', 'Khó copy chính xác'),
        ('Font chữ', '15%', 'Brand identity'),
        ('Mã QR', '15%', 'Phải khớp'),
        ('Màu sắc', '10%', 'Có sai số do ánh sáng'),
    ],
    col_widths=[5, 3, 9]
)

add_heading('5.7 Hybrid Review Workflow', 2)
p = doc.add_paragraph()
run = p.add_run('''
AI phân tích
      ↓
Risk Score
      ↓
< 40   →  Tự động thông qua (auto pass)
40-70  →  Chuyển kiểm tra nhân công (manual review)
> 70   →  Cảnh báo ĐỎ → Điều tra thực địa (field investigation)
''')
run.font.name = 'Consolas'
run.font.size = Pt(10)

add_heading('5.8 Investigation Console (MỚI — chưa có)', 2, RED)
add_bullet('Timeline vụ việc với từng action (upload → AI alert → assign QLTT → kết luận)')
add_bullet('Side-by-side image comparison (ảnh nghi vấn vs ảnh gốc) + diff %')
add_bullet('Notes, evidence attachments, status changes')
add_bullet('Assign cán bộ điều tra + due date + escalation')
add_bullet('Export báo cáo vụ việc PDF')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# VI. MODULE 4 — BI ANALYTICS
# ═══════════════════════════════════════════════════════════════════════
add_heading('VI. MODULE 4 — BI/ANALYTICS COMMAND CENTER', 1, GOLD)

add_heading('6.1 12 mục menu Command Center', 2)
items = [
    '1. Dashboard tổng quan', '2. Tra cứu nhanh (Smart Search)',
    '3. Phân tích hàng giả', '4. Phân tích doanh nghiệp',
    '5. Heatmap GPS', '6. Theo dõi thương hiệu',
    '7. AI Risk Analytics', '8. Điều tra & vụ việc',
    '9. Báo cáo thống kê', '10. Cảnh báo thời gian thực',
    '11. Xuất dữ liệu / API', '12. Quản trị hệ thống',
]
for it in items:
    add_bullet(it)

add_heading('6.2 Smart Search đa chiều', 2)
add_table(
    ['Nhóm', 'Trường tra cứu'],
    [
        ('Doanh nghiệp', 'Tên DN, MST, Địa chỉ'),
        ('Sản phẩm', 'SKU, Barcode, QR, Batch'),
        ('Thương hiệu', 'Logo, Nhãn hiệu'),
        ('Vi phạm', 'Hàng giả, Gian lận, Kém chất lượng'),
        ('Địa lý', 'Tỉnh, Quận, GPS bbox'),
        ('Thời gian', 'Ngày, Tuần, Tháng, Quý, Năm'),
    ],
    col_widths=[4, 13]
)

add_heading('6.3 8 loại biểu đồ', 2)
add_table(
    ['Loại', 'Mục tiêu', 'Use case'],
    [
        ('Bar Chart', 'So sánh giá trị', 'Top 10 SP bị giả mạo'),
        ('Line Chart', 'Xu hướng theo time', 'Số scan/ngày'),
        ('Pie Chart', 'Phân loại tỷ lệ', 'Cơ cấu vi phạm theo loại'),
        ('Heatmap GIS', 'Mật độ địa lý', 'Bản đồ điểm nóng hàng giả'),
        ('Sankey', 'Dòng phân phối', 'Lô hàng từ NSX → NPP → retail'),
        ('Network Graph', 'Liên kết DN', 'Mạng lưới shop bán hàng giả'),
        ('Funnel', 'Conversion', 'Quét → báo cáo → xử lý'),
        ('Treemap', 'Hierarchy', 'Cơ cấu ngành/sub-ngành'),
    ],
    col_widths=[3, 5, 8]
)

add_heading('6.4 Time slider animation', 2)
add_bullet('Slider chia thành 24h / 7d / 30d buckets')
add_bullet('Auto-play 4 tốc độ (×½, ×1, ×2, ×3)')
add_bullet('Heatmap evolve theo thời gian → phát hiện xu hướng lan rộng')
add_bullet('Đã có: ✅ VietMap Phase 4 (Apr 2026)')

add_heading('6.5 GIS Heatmap 5 màu (chuẩn hóa)', 2)
add_table(
    ['Màu', 'Ý nghĩa', 'Action'],
    [
        ('🟢 Xanh lá', 'An toàn (<10 cases)', 'Không cần can thiệp'),
        ('🟡 Vàng', 'Theo dõi (10-50)', 'Monitor weekly'),
        ('🟠 Cam', 'Nghi vấn (50-100)', 'Cảnh báo QLTT'),
        ('🔴 Đỏ', 'Điểm nóng (>100)', 'Điều tra tức thì'),
        ('⚫ Đen', 'Khu vực trọng điểm', 'Mobile QLTT thường trực'),
    ],
    col_widths=[4, 5, 8]
)

add_heading('6.6 Realtime Dashboard widgets', 2)
add_bullet('⚠ N cảnh báo mới (cập nhật mỗi 30s)')
add_bullet('⚠ M QR clone phát hiện 24h gần đây')
add_bullet('⚠ K thương hiệu bị giả mạo tuần này')
add_bullet('⚠ L sản phẩm hết hạn trong tuần')
add_bullet('⚠ X điểm nóng mới trên map')

add_heading('6.7 Export formats', 2)
add_table(
    ['Định dạng', 'Mục tiêu', 'Status'],
    [
        ('PDF', 'Báo cáo điều hành', '✅'),
        ('Excel', 'Phân tích chi tiết', '✅'),
        ('CSV', 'Data raw export', '✅'),
        ('JSON API', 'Liên thông hệ thống', '✅'),
        ('Power BI / Tableau', 'BI nâng cao', '❌ Chưa có connector'),
    ],
    col_widths=[5, 7, 4]
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# VII. PHÂN QUYỀN + KYC
# ═══════════════════════════════════════════════════════════════════════
add_heading('VII. PHÂN QUYỀN, KYC & SUBSCRIPTION', 1, GOLD)

add_heading('7.1 4 platform roles', 2)
add_table(
    ['Role', 'Mô tả', 'Cách tạo'],
    [
        ('admin', 'Quản trị hệ thống — toàn quyền', 'Seed DB (không UI)'),
        ('manufacturer', 'NSX — quản lý SP/lô/QR', 'Đăng ký /login/manufacturer'),
        ('importer', 'NNK — quản lý nhập khẩu', 'Đăng ký /login/importer'),
        ('consumer', 'Người tiêu dùng', 'Đăng ký /login/consumer hoặc guest'),
    ],
    col_widths=[4, 8, 5]
)

add_heading('7.2 4 sub-role nội bộ DN (UC03)', 2)
add_table(
    ['Sub-role', 'Quyền'],
    [
        ('company_admin', 'Toàn quyền + mời nhân viên'),
        ('staff_input', 'Nhập SP/lô + tạo QR'),
        ('warehouse', 'Ghi nhập/xuất kho'),
        ('viewer', 'Chỉ xem dashboard + báo cáo'),
    ],
    col_widths=[5, 12]
)

add_heading('7.3 Roles đặc biệt (MỚI — chưa có)', 2, RED)
add_bullet('🔴 Cán bộ tư vấn — UC19 GTIN check + lịch sử tra cứu cross-DN')
add_bullet('🔴 Cơ quan chức năng — QLTT/Hải quan/ATTP với read-only cross-DN + unmask flow')
add_bullet('🔴 Thanh tra viên — Investigation Console + assign vụ việc')

add_heading('7.4 KYC workflow', 2)
add_number('DN đăng ký tại /login/manufacturer hoặc /login/importer')
add_number('Hệ thống tạo account trạng thái pending (chờ duyệt)')
add_number('Admin vào /dashboard/kyc kiểm tra hồ sơ + chứng từ')
add_number('Admin duyệt (verified) hoặc từ chối (suspended)')
add_number('Sau khi verified, DN mới được đăng nhập + cập nhật dữ liệu')

add_heading('7.5 4 gói Subscription (MỚI — chưa có Payment system)', 2, RED)
add_table(
    ['Gói', 'Quota QR/tháng', 'Tính năng', 'API'],
    [
        ('Basic', '1,000', 'Hồ sơ + lô + scan', 'Không'),
        ('Pro', '10,000', 'Basic + Webhook + Heatmap', 'REST'),
        ('Enterprise', '100,000', 'Pro + AI Dashboard + IOC', 'REST + Webhook'),
        ('Government', 'Unlimited', 'All + ưu tiên + audit chi tiết', 'Custom'),
    ],
    col_widths=[4, 4, 7, 3]
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# VIII. BẢO MẬT
# ═══════════════════════════════════════════════════════════════════════
add_heading('VIII. BẢO MẬT, VAULT & PHÁP LÝ', 1, GOLD)

add_heading('8.1 Identity Vault 3 kho tách biệt', 2)
add_table(
    ['Vault', 'Lưu', 'Khóa mã'],
    [
        ('IdentityVault', 'Hash(UserID) → ReportID (không lưu raw)', 'AES-256'),
        ('ReportVault', 'Nội dung báo cáo nghi vấn', 'AES-256 + per-DN key'),
        ('AnonymousSession', 'Phiên ẩn danh tạm thời (TTL 24h)', 'Random token'),
    ],
    col_widths=[5, 8, 4]
)

add_heading('8.2 Quy trình unmask khi cơ quan chức năng yêu cầu', 2, RED)
add_number('Cơ quan gửi yêu cầu kèm văn bản chính thức (PDF có chữ ký số)')
add_number('Admin upload văn bản → tạo ticket UnmaskRequest trên hệ thống')
add_number('Pháp chế DN review → approve/reject')
add_number('Nếu approve: hệ thống decrypt mapping Hash → User thật → log audit')
add_number('Trả PII cho cơ quan qua channel bảo mật (PGP encrypted email)')

add_heading('8.3 8 biện pháp bảo mật API', 2)
add_bullet('OAuth2 + JWT HS256 (TTL 30 phút)')
add_bullet('CSP whitelist domain nghiêm ngặt')
add_bullet('Rate limit per IP + per user')
add_bullet('HMAC-SHA256 cho webhook outbound')
add_bullet('PKI chữ ký số cho documents quan trọng')
add_bullet('mTLS internal cho microservices')
add_bullet('SIEM log centralized')
add_bullet('AES-256 mã hóa data at rest')

add_heading('8.4 Tuân thủ pháp lý', 2)
add_bullet('Luật Bảo vệ dữ liệu cá nhân 2025 (hiệu lực 1/1/2026)')
add_bullet('Luật An ninh mạng')
add_bullet('Luật Dữ liệu 2025 (hiệu lực 1/7/2025)')
add_bullet('NĐ 37/2026 chống hàng giả')
add_bullet('TCVN 13275:2020 truy xuất nguồn gốc')
add_bullet('ISO 27001 ATTT')
add_bullet('Khung EA Việt Nam (Chính phủ điện tử)')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# IX. USE CASES
# ═══════════════════════════════════════════════════════════════════════
add_heading('IX. USE CASES (UC01–UC30)', 1, GOLD)

uc_list = [
    ('UC01', 'Đăng ký Tài khoản DN', '✅'),
    ('UC02', 'Upload + Xác thực Chứng từ (KYC)', '✅'),
    ('UC03', 'Quản lý Phân quyền Nội bộ DN (4 sub-role)', '✅'),
    ('UC04', 'Thêm/Sửa/Xóa Sản phẩm', '✅'),
    ('UC05', 'Quản lý Lô hàng', '✅'),
    ('UC06', 'Đính kèm Chứng nhận', '✅'),
    ('UC07', 'Tạo Mã UID & QR Hàng loạt', '✅'),
    ('UC08', 'Xuất File In ấn Mã QR', '✅'),
    ('UC09', 'Quét QR Xác thực Sản phẩm', '✅'),
    ('UC10', 'Xem Thông tin Chi tiết SP', '✅'),
    ('UC11', 'Báo cáo Nghi ngờ Hàng giả', '✅'),
    ('UC12', 'Xem Dashboard Giám sát', '✅'),
    ('UC13', 'Xem Bản đồ Quét Địa lý (Heatmap)', '✅'),
    ('UC14', 'Cấu hình Ngưỡng Cảnh báo', '✅'),
    ('UC15', 'Xem & Xử lý Cảnh báo', '✅'),
    ('UC16', 'Xuất Báo cáo Tóm tắt', '✅'),
    ('UC17', 'Xuất Báo cáo Chi tiết Lô hàng', '✅'),
    ('UC18', 'Kiểm tra Tuân thủ Chứng nhận (TrustCheck)', '✅'),
    ('UC19', 'Kiểm tra Tính hợp lệ GTIN', '✅'),
    ('UC20', 'Quản lý Người dùng (Admin)', '✅'),
    ('UC21 (mới)', 'AI Photo Vision Verify', '✅'),
    ('UC22 (mới)', 'Cấu hình Webhook ERP', '✅'),
    ('UC23 (mới)', 'Cấu hình Vòng đời theo DN', '✅'),
    ('UC24 (mới)', 'Geocoding DN address', '✅'),
    ('UC25 (mới)', 'Time slider heatmap animation', '✅'),
    ('UC26 (cần)', 'Investigation Console (timeline + assign)', '❌'),
    ('UC27 (cần)', 'Side-by-side Image Comparison + diff %', '❌'),
    ('UC28 (cần)', 'Reward System cho NTD', '❌'),
    ('UC29 (cần)', 'Subscription Payment (Basic/Pro/Ent/Gov)', '❌'),
    ('UC30 (cần)', 'Unmask PII flow cho cơ quan chức năng', '❌'),
]
add_table(
    ['Code', 'Tên Use Case', 'Status'],
    uc_list,
    col_widths=[3, 11, 3]
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# X. CẢNH BÁO VÒNG ĐỜI + HẬU KIỂM
# ═══════════════════════════════════════════════════════════════════════
add_heading('X. CẢNH BÁO VÒNG ĐỜI & HẬU KIỂM', 1, GOLD)

add_heading('10.1 3 lớp cảnh báo vòng đời', 2)
add_bullet('LỚP 1 — Chứng nhận DN (ISO, HACCP, Halal) hết hạn → ngày: 90/30/7')
add_bullet('LỚP 2 — Hạn sử dụng sản phẩm (EXP) → tùy ngành: Sữa 60d, Mỹ phẩm 90d, Dược 30d')
add_bullet('LỚP 3 — Phiên bản tiêu chuẩn lỗi thời')

add_heading('10.2 Kênh thông báo đa kênh', 2)
add_bullet('Email (Gmail SMTP — đã có)')
add_bullet('SMS (Twilio — chưa có)')
add_bullet('Push notification (Web + Mobile)')
add_bullet('Webhook ERP (HMAC-SHA256 signed)')
add_bullet('Zalo OA (giai đoạn 2)')

add_heading('10.3 Hậu kiểm — 3 đối tượng upload', 2)
add_bullet('Doanh nghiệp sản xuất tự lấy mẫu')
add_bullet('Người tiêu dùng thuê labs phân tích')
add_bullet('Cơ quan QLTT / Công an / ATTP lấy mẫu thanh tra')

add_heading('10.4 Rule Engine 25+ tiêu chuẩn', 2)
add_bullet('Auto compare value với max/min/range/qualitative thresholds')
add_bullet('Kết luận: ĐẢM BẢO / KHÔNG ĐẢM BẢO')
add_bullet('Audit trail rule engine version')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# XI. STATUS HIỆN TẠI + GAP ANALYSIS
# ═══════════════════════════════════════════════════════════════════════
add_heading('XI. TRẠNG THÁI HIỆN TẠI & GAP ANALYSIS', 1, GOLD)

add_heading('11.1 Mức tuân thủ theo module', 2)
add_table(
    ['Module', 'Compliance', 'Status'],
    [
        ('Module 1 — Enterprise Portal', '85%', '🟢 Tốt, thiếu Nhà máy/Logo/ESG model'),
        ('Module 2 — Consumer Crowdsourcing', '75%', '🟡 Thiếu Wizard, Reward, Livestream'),
        ('Module 3 — AI Fraud Detection', '60%', '🟠 Thiếu Investigation Console, Side-by-side, YOLOv8'),
        ('Module 4 — BI/Analytics', '80%', '🟢 Có Heatmap + Time slider, thiếu Sankey/Network'),
        ('Module 5 — Unified Portal', '40%', '🔴 Thiếu Landing 4 nút + Payment'),
        ('Phân quyền + KYC + UC03', '95%', '🟢 Đầy đủ'),
        ('Bảo mật + Vault', '85%', '🟢 Có vault, thiếu Unmask UI'),
        ('Cảnh báo vòng đời + Hậu kiểm', '95%', '🟢 Đầy đủ'),
        ('TỔNG THỂ', '~83%', '🟢 Production-ready 80%+'),
    ],
    col_widths=[6, 4, 7]
)

add_heading('11.2 Gap urgent — cần làm ngay', 2)
add_bullet('🔴 Landing page Unified Portal 4 nút lớn (file 5)')
add_bullet('🔴 Investigation Console (file 3+4)')
add_bullet('🔴 Checklist UI 6-color severity (file 3)')
add_bullet('🔴 Side-by-side image comparison (file 3)')
add_bullet('🔴 Subscription Payment system (file 5)')
add_bullet('🔴 Unmask PII flow (file 3 — §III.9)')

add_heading('11.3 Gap medium — sprint sau', 2)
add_bullet('🟡 Nhà máy + Thương hiệu + ESG models (file 1)')
add_bullet('🟡 Wizard 4-bước upload (file 2)')
add_bullet('🟡 Reward gamification (file 2)')
add_bullet('🟡 Cán bộ tư vấn + Cơ quan chức năng roles')
add_bullet('🟡 YOLOv8 + PaddleOCR + Siamese CNN server-side')
add_bullet('🟡 Sankey + Network Graph charts')

add_heading('11.4 Gap low — future', 2)
add_bullet('🟢 Mobile native Flutter')
add_bullet('🟢 Livestream Detection AI')
add_bullet('🟢 Social Media Monitoring')
add_bullet('🟢 Blockchain traceability')
add_bullet('🟢 Edge AI offline')
add_bullet('🟢 Camunda BPMN workflow engine')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# XII. ROADMAP
# ═══════════════════════════════════════════════════════════════════════
add_heading('XII. ROADMAP 3 GIAI ĐOẠN', 1, GOLD)

add_heading('Giai đoạn 1 — MVP (Đã hoàn thành ~85%)', 2, GREEN)
add_bullet('✅ Enterprise Portal cơ bản + KYC')
add_bullet('✅ Consumer scan + report')
add_bullet('✅ AI verify image basic')
add_bullet('✅ Heatmap 4 layers + Time slider')
add_bullet('✅ Webhook ERP + Lifecycle config + System config')

add_heading('Giai đoạn 2 — Nâng cao (Q3/2026)', 2, AMBER)
add_bullet('🟡 Landing Unified Portal 4 nút')
add_bullet('🟡 Investigation Console + 6-color checklist')
add_bullet('🟡 Subscription Payment')
add_bullet('🟡 Wizard Consumer + Reward')
add_bullet('🟡 Unmask PII flow')
add_bullet('🟡 Server-side AI (YOLOv8 + PaddleOCR)')
add_bullet('🟡 Mobile PWA installable')

add_heading('Giai đoạn 3 — Mở rộng (Q1/2027)', 2, BLUE)
add_bullet('🔵 Mobile native Flutter')
add_bullet('🔵 Livestream Detection')
add_bullet('🔵 Social Media Monitoring')
add_bullet('🔵 Blockchain traceability')
add_bullet('🔵 Edge AI offline')
add_bullet('🔵 National IOC dashboard')
add_bullet('🔵 Integration QLTT/Hải quan/ATTP nationwide')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════════════
add_heading('PHỤ LỤC', 1, GOLD)

add_heading('A. Tech Stack hiện tại', 2)
add_table(
    ['Component', 'Technology', 'Recommended (file 1-5)'],
    [
        ('Frontend', 'Next.js 16 + React 19 + Tailwind', 'React/NextJS ✅'),
        ('Backend API', 'Next.js API Routes', 'NestJS ❌'),
        ('Database', 'SQLite (dev), PostgreSQL (prod)', 'PostgreSQL ✅'),
        ('ORM', 'Prisma 7', '—'),
        ('Auth', 'JWT HS256 + bcrypt', 'OAuth2 + SSO 🟡'),
        ('Map', 'MapLibre GL + CartoDB + Esri', 'Mapbox ≈'),
        ('AI Client', 'TF.js MobileNet (CDN)', 'TensorFlow ✅'),
        ('AI Server', 'Không có', 'YOLOv8 + Siamese CNN ❌'),
        ('OCR', 'Tesseract.js stub', 'PaddleOCR ❌'),
        ('Workflow', 'Custom Prisma logic', 'Camunda ❌'),
        ('Search', 'Prisma queries', 'Elasticsearch + Vector ❌'),
        ('Stream', 'Không có', 'Kafka ❌'),
        ('Log analytics', 'NhatKy table', 'ClickHouse ❌'),
        ('BI', 'Custom dashboards', 'Apache Superset ❌'),
        ('Email', 'Nodemailer + Gmail SMTP', 'SendGrid/SES ≈'),
    ],
    col_widths=[4, 7, 6]
)

add_heading('B. Schema Prisma hiện tại — 21 models', 2)
models = [
    'DoanhNghiep', 'SanPham', 'LoHang', 'KhoHang', 'MaDinhDanh', 'LuotQuet',
    'CanhBao', 'ChungNhan', 'NguoiDung', 'LoiMoiNhanVien', 'KetQuaHauKiem',
    'CauHinhHeThong', 'YeuCauTuanThuVSIC', 'IdentityVault', 'ReportVault',
    'AnonymousSession', 'WebhookErp', 'TieuChuanKiemNghiem', 'NhatKy', 'ThongBao',
    'DonChuyenHang',
]
for m in models:
    add_bullet(m)

add_heading('C. Models CẦN BỔ SUNG (theo 5 file mới)', 2, RED)
add_table(
    ['Model mới', 'Mục đích'],
    [
        ('NhaMay', 'Thông tin nhà máy + dây chuyền + SOP'),
        ('ThuongHieu', 'Logo + nhãn hiệu + bảo hộ SHTT'),
        ('Investigation / Case', 'Timeline vụ việc + assign cán bộ'),
        ('UserScore / Reward', 'Gamification cho NTD'),
        ('Subscription', 'Gói phí Basic/Pro/Enterprise/Gov'),
        ('Payment', 'Lịch sử thanh toán + invoice'),
        ('SocialMention', 'Monitoring MXH Facebook/TikTok'),
        ('UnmaskRequest', 'Yêu cầu giải mã PII từ cơ quan'),
        ('ESGReport', 'Báo cáo ESG/CSR của DN'),
    ],
    col_widths=[5, 12]
)

add_heading('D. 6 tài liệu nguồn', 2)
sources = [
    'TAI_LIEU_NGHIEP_VU_VNTRUST_FULL.docx — Tài liệu chính 10 chương',
    '1. hệ thống dữ liệu sản phẩm doanh nghiệp.docx — Module 1',
    '2. hệ thống phản ánh sản phẩm hàng hóa từ khách hàng.docx — Module 2',
    '3. hệ thống chứa các thuật toán để kiểm tra.docx — Module 3',
    '4. hệ thống truy vấn dữ liệu và kết xuất báo cáo.docx — Module 4',
    '5. Giao diện chung của toàn bộ nền tảng.docx — Module 5 / Portal',
]
for s in sources:
    add_bullet(s)

# Footer
doc.add_paragraph()
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot_run = foot.add_run('— HẾT TÀI LIỆU V3.0 —')
foot_run.italic = True
foot_run.font.color.rgb = GOLD
foot_run.font.size = Pt(11)

# ─── Save ────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'✅ Generated: {OUT}')
print(f'   Size: {os.path.getsize(OUT)/1024:.1f} KB')
print(f'   Paragraphs: {len(doc.paragraphs)}')
print(f'   Tables: {len(doc.tables)}')
